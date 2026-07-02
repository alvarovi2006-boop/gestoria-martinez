#!/usr/bin/env python3
"""
generar_informes.py

Lee un archivo Excel con las columnas: fecha, concepto, importe, categoría
(ingreso/gasto) y genera un documento Word con:

  1. Una tabla resumen de ingresos y gastos totales del mes.
  2. Una gráfica de barras comparando ingresos vs gastos.

Uso:
    python generar_informes.py [archivo_entrada.xlsx] [informe_salida.docx]

Si no se indican argumentos usa 'cliente_ejemplo.xlsx' como entrada y
'informe.docx' como salida.

Dependencias:
    pip install pandas openpyxl matplotlib python-docx
"""

import sys
import unicodedata
from pathlib import Path

import pandas as pd
import matplotlib

matplotlib.use("Agg")  # backend sin ventana, para generar imágenes en disco
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Utilidades
# ---------------------------------------------------------------------------

MESES_ES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
}


def _normaliza(texto: str) -> str:
    """Pasa a minúsculas y quita acentos, para comparar nombres de columna."""
    texto = str(texto).strip().lower()
    texto = unicodedata.normalize("NFKD", texto)
    return "".join(c for c in texto if not unicodedata.combining(c))


def cargar_datos(ruta: Path) -> pd.DataFrame:
    """Carga el Excel y normaliza las columnas esperadas."""
    if not ruta.exists():
        raise FileNotFoundError(f"No se encuentra el archivo de entrada: {ruta}")

    df = pd.read_excel(ruta)

    # Mapea los nombres de columna reales a los nombres canónicos que usamos.
    alias = {
        "fecha": "fecha",
        "concepto": "concepto",
        "importe": "importe",
        "categoria": "categoria",  # 'categoría' se normaliza a 'categoria'
    }
    columnas = {}
    for col in df.columns:
        clave = _normaliza(col)
        if clave in alias:
            columnas[col] = alias[clave]
    df = df.rename(columns=columnas)

    requeridas = {"fecha", "concepto", "importe", "categoria"}
    faltan = requeridas - set(df.columns)
    if faltan:
        raise ValueError(
            "Faltan columnas obligatorias en el Excel: "
            + ", ".join(sorted(faltan))
            + f"\nColumnas encontradas: {list(df.columns)}"
        )

    # Tipos y limpieza.
    df["fecha"] = pd.to_datetime(df["fecha"], errors="coerce")
    df["importe"] = pd.to_numeric(df["importe"], errors="coerce").fillna(0.0)
    df["categoria"] = df["categoria"].apply(_normaliza)

    # Nos quedamos solo con ingreso/gasto reconocidos.
    validas = df["categoria"].isin(["ingreso", "gasto"])
    if not validas.all():
        desconocidas = df.loc[~validas, "categoria"].unique()
        print(f"Aviso: se ignoran filas con categoría desconocida: {desconocidas}")
    df = df[validas].copy()

    if df.empty:
        raise ValueError("No hay transacciones válidas (ingreso/gasto) que procesar.")

    return df


def titulo_periodo(df: pd.DataFrame) -> str:
    """Devuelve un texto tipo 'julio de 2026' a partir del mes predominante."""
    fechas = df["fecha"].dropna()
    if fechas.empty:
        return "periodo sin fecha"
    # Mes/año más frecuente en los datos.
    periodo = fechas.dt.to_period("M").mode().iloc[0]
    return f"{MESES_ES[periodo.month]} de {periodo.year}"


# ---------------------------------------------------------------------------
# Gráfica
# ---------------------------------------------------------------------------

def generar_grafica(total_ingresos: float, total_gastos: float, ruta_img: Path) -> None:
    """Crea una gráfica de barras comparando ingresos vs gastos."""
    categorias = ["Ingresos", "Gastos"]
    valores = [total_ingresos, total_gastos]
    colores = ["#2e7d32", "#c62828"]  # verde ingresos, rojo gastos

    fig, ax = plt.subplots(figsize=(6, 4))
    barras = ax.bar(categorias, valores, color=colores, width=0.55)

    ax.set_title("Ingresos vs Gastos", fontsize=14, fontweight="bold")
    ax.set_ylabel("Importe (€)")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", linestyle="--", alpha=0.4)

    # Etiqueta con el valor encima de cada barra.
    for barra, valor in zip(barras, valores):
        ax.annotate(
            f"{valor:,.2f} €",
            xy=(barra.get_x() + barra.get_width() / 2, valor),
            xytext=(0, 4),
            textcoords="offset points",
            ha="center",
            va="bottom",
            fontsize=10,
            fontweight="bold",
        )

    fig.tight_layout()
    fig.savefig(ruta_img, dpi=150)
    plt.close(fig)


# ---------------------------------------------------------------------------
# Documento Word
# ---------------------------------------------------------------------------

def generar_word(df: pd.DataFrame, ruta_img: Path, ruta_docx: Path) -> None:
    total_ingresos = df.loc[df["categoria"] == "ingreso", "importe"].sum()
    total_gastos = df.loc[df["categoria"] == "gasto", "importe"].sum()
    balance = total_ingresos - total_gastos
    periodo = titulo_periodo(df)

    doc = Document()

    # --- Título ---
    titulo = doc.add_heading("Informe económico", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Periodo: {periodo}")
    run.italic = True
    run.font.size = Pt(12)

    # --- Tabla resumen ---
    doc.add_heading("Resumen de ingresos y gastos", level=1)

    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Light Grid Accent 1"

    hdr = tabla.rows[0].cells
    hdr[0].text = "Concepto"
    hdr[1].text = "Importe (€)"
    for celda in hdr:
        for parrafo in celda.paragraphs:
            for r in parrafo.runs:
                r.bold = True

    filas = [
        ("Total ingresos", total_ingresos),
        ("Total gastos", total_gastos),
        ("Balance", balance),
    ]
    for concepto, importe in filas:
        celdas = tabla.add_row().cells
        celdas[0].text = concepto
        celdas[1].text = f"{importe:,.2f}"
        celdas[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Colorea el balance según sea positivo o negativo.
    fila_balance = tabla.rows[-1].cells
    color = RGBColor(0x2E, 0x7D, 0x32) if balance >= 0 else RGBColor(0xC6, 0x28, 0x28)
    for celda in fila_balance:
        for parrafo in celda.paragraphs:
            for r in parrafo.runs:
                r.bold = True
                r.font.color.rgb = color

    doc.add_paragraph()

    # --- Gráfica ---
    doc.add_heading("Comparativa gráfica", level=1)
    doc.add_picture(str(ruta_img), width=Inches(5.5))
    ultimo = doc.paragraphs[-1]
    ultimo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Detalle de transacciones ---
    doc.add_heading("Detalle de transacciones", level=1)
    detalle = doc.add_table(rows=1, cols=4)
    detalle.style = "Light List Accent 1"
    cab = detalle.rows[0].cells
    for i, texto in enumerate(["Fecha", "Concepto", "Categoría", "Importe (€)"]):
        cab[i].text = texto
        for parrafo in cab[i].paragraphs:
            for r in parrafo.runs:
                r.bold = True

    for _, fila in df.sort_values("fecha").iterrows():
        celdas = detalle.add_row().cells
        fecha_txt = fila["fecha"].strftime("%d/%m/%Y") if pd.notna(fila["fecha"]) else "-"
        celdas[0].text = fecha_txt
        celdas[1].text = str(fila["concepto"])
        celdas[2].text = str(fila["categoria"]).capitalize()
        celdas[3].text = f"{fila['importe']:,.2f}"
        celdas[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(ruta_docx)


# ---------------------------------------------------------------------------
# Programa principal
# ---------------------------------------------------------------------------

def main() -> None:
    entrada = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("cliente_ejemplo.xlsx")
    salida = Path(sys.argv[2]) if len(sys.argv) > 2 else Path("informe.docx")
    ruta_img = salida.with_name("grafica_ingresos_gastos.png")

    print(f"Leyendo datos de: {entrada}")
    df = cargar_datos(entrada)

    total_ingresos = df.loc[df["categoria"] == "ingreso", "importe"].sum()
    total_gastos = df.loc[df["categoria"] == "gasto", "importe"].sum()

    print(f"  Total ingresos: {total_ingresos:,.2f} €")
    print(f"  Total gastos:   {total_gastos:,.2f} €")
    print(f"  Balance:        {total_ingresos - total_gastos:,.2f} €")

    print("Generando gráfica...")
    generar_grafica(total_ingresos, total_gastos, ruta_img)

    print("Generando documento Word...")
    generar_word(df, ruta_img, salida)

    print(f"\n¡Listo! Informe guardado en: {salida.resolve()}")
    print(f"Gráfica guardada en:       {ruta_img.resolve()}")


if __name__ == "__main__":
    main()
